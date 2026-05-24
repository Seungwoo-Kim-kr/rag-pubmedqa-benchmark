"""
make_report.py  —  supervisor 보고서 생성
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 설정 (A4, 여백 2.5cm) ──────────────────────────────
section = doc.sections[0]
section.page_width  = int(21.0 * 914400 / 25.4)
section.page_height = int(29.7 * 914400 / 25.4)
for attr in ("top_margin","bottom_margin","left_margin","right_margin"):
    setattr(section, attr, Cm(2.5))

# ── 헬퍼 ──────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(31, 73, 125) if level==1 else RGBColor(54, 95, 145)
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5) if indent else Cm(0)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(10.5)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text)
        set_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_font(r, size=10.5)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_table(doc, headers, rows, col_widths, header_color="1F497D"):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade_cell(cell, header_color)
        cell.width = Cm(col_widths[j])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9.5, bold=True, color=(255,255,255))
    # 데이터
    for i, row in enumerate(rows):
        bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            shade_cell(cell, bg)
            cell.width = Cm(col_widths[j])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            bold = j == 0
            set_font(run, size=9.5, bold=bold)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("RAG 기반 과학 문서 QA 비교 실험\n연구 진행 보고서")
set_font(r, size=20, bold=True, color=(31,73,125))

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("2026년 5월")
set_font(r2, size=12, color=(89,89,89))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. 프로젝트 개요
# ══════════════════════════════════════════════════════════════
heading(doc, "1. 프로젝트 개요")
body(doc,
    "본 실험은 과학 문서를 대상으로 4가지 RAG(Retrieval-Augmented Generation) "
    "방식의 QA 성능을 체계적으로 비교하기 위해 설계되었습니다. "
    "단순 성능 측정에 그치지 않고, 기존 평가 지표(Token F1)의 한계를 검증하고 "
    "LLM-as-judge 방식의 보완 필요성을 실험적으로 입증하는 것을 목표로 합니다.")

heading(doc, "실험 목표", level=2)
bullet(doc, "4가지 baseline 방식의 정량적 성능 비교")
bullet(doc, "질문 유형(local_factual / global_synthesis / terminology_sensitive)별 강약점 분석")
bullet(doc, "Token F1 vs. LLM Judge 평가 지표 간 괴리 측정")
bullet(doc, "Human-annotated 데이터셋(QASPER)과 자동 생성 데이터셋(PubMedQA) 간 평가 난이도 비교")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 2. 파이프라인 구성
# ══════════════════════════════════════════════════════════════
heading(doc, "2. 실험 파이프라인 구성")

heading(doc, "2-1. 데이터 준비", level=2)
add_table(doc,
    ["구분", "데이터셋", "질문 수", "문서 수", "질문 출처"],
    [
        ["데이터셋 A", "PubMedQA", "45개 (타입별 15)", "50개", "GPT-4o-mini 자동 생성"],
        ["데이터셋 B", "QASPER (NLP 논문)", "45개 (타입별 15)", "52개", "Human-annotated"],
    ],
    col_widths=[2.5, 3.2, 3.0, 2.0, 4.5]
)
bullet(doc, "질문 타입 균형: local_factual 15개 / global_synthesis 15개 / terminology_sensitive 15개")
bullet(doc, "QASPER: AllenAI 공개 train set (888개 논문) 중 파싱, BIBREF/FIGREF 등 참조 태그 제거")
bullet(doc, "PubMedQA: HuggingFace API를 통해 직접 로드, GPT로 균형 질문 생성")

doc.add_paragraph()
heading(doc, "2-2. Baseline 4종", level=2)
add_table(doc,
    ["Baseline", "방식 설명", "LLM 호출 수", "특이사항"],
    [
        ["Direct QA",          "문서 전체를 컨텍스트로 LLM에 직접 전달",     "1회", "정보 손실 없음, 장문 문서 시 컨텍스트 초과 위험"],
        ["Standard RAG",       "FAISS로 top-k 청크 검색 후 LLM에 전달",       "1회", "빠름, terminology 질문에 취약"],
        ["Summary-Mediated QA","청크 검색 → 요약 → 답변 생성 2단계",          "2회", "비용·시간 2배, 가장 안정적"],
        ["LightRAG (hybrid)",  "지식 그래프(개체-관계) 기반 하이브리드 검색", "복수","Python 3.12 별도 환경 필요"],
    ],
    col_widths=[3.5, 5.0, 2.2, 5.0]
)

heading(doc, "2-3. 평가 지표", level=2)
add_table(doc,
    ["지표", "설명", "한계"],
    [
        ["Token F1",     "Gold 답변과 예측 답변의 토큰 겹침 F1 (SQuAD 방식)", "동의어·paraphrase 시 과소평가"],
        ["Retrieval Hit","Gold evidence가 retrieved context에 50%+ 포함 여부", "Binary 측정, 세밀도 부족"],
        ["Judge Score",  "GPT-4o-mini가 correct/partial/incorrect 3단계 채점", "동일 모델 채점으로 편향 가능성"],
    ],
    col_widths=[3.0, 7.0, 5.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 3. 실험 결과
# ══════════════════════════════════════════════════════════════
heading(doc, "3. 실험 결과")

heading(doc, "3-1. PubMedQA — GPT 생성 질문 (45문항 × 4 baseline)", level=2)
add_table(doc,
    ["Baseline", "Token F1", "Judge Score", "Hit Rate", "Correct", "Partial", "Incorrect", "응답시간"],
    [
        ["Summary-Mediated QA", "0.526", "0.878 ★", "1.00", "34", "11", "0",  "4.1s"],
        ["Direct QA",           "0.497", "0.844",   "0.49", "33", "10", "2",  "3.2s"],
        ["Standard RAG",        "0.364", "0.622",   "1.00", "23", "10", "12", "1.9s ★"],
        ["LightRAG (hybrid)",   "0.150", "0.278",   "0.00", "3",  "19", "23", "5.6s"],
    ],
    col_widths=[3.8, 2.0, 2.5, 2.0, 1.8, 1.8, 2.0, 2.0]
)

heading(doc, "질문 타입별 Judge Score (PubMedQA)", level=2)
add_table(doc,
    ["Baseline", "local_factual", "global_synthesis", "terminology_sensitive"],
    [
        ["Summary-Mediated QA", "0.967", "0.867", "0.800 ★"],
        ["Direct QA",           "0.967", "0.867", "0.700"],
        ["Standard RAG",        "0.900", "0.767", "0.200 ▼"],
        ["LightRAG (hybrid)",   "0.067 ▼","0.400","0.367"],
    ],
    col_widths=[4.0, 3.5, 3.8, 4.2]
)

body(doc, "▶ 핵심 패턴: Standard RAG는 terminology_sensitive에서 Judge=0.200으로 급락. "
         "LightRAG는 local_factual에서 0.067로 단일 문서 조건에서의 한계 노출.", indent=True)

doc.add_paragraph()
heading(doc, "3-2. QASPER — Human-annotated 질문 (45문항 × 3 baseline)", level=2)
add_table(doc,
    ["Baseline", "Token F1", "Judge Score", "Hit Rate", "Correct", "Partial", "Incorrect", "응답시간"],
    [
        ["Direct QA",           "0.240", "0.444 ★", "0.04", "10", "20", "15", "2.3s"],
        ["Summary-Mediated QA", "0.222", "0.400",   "1.00", "5",  "26", "14", "4.2s"],
        ["Standard RAG",        "0.177", "0.322",   "1.00", "5",  "19", "21", "2.3s"],
    ],
    col_widths=[3.8, 2.0, 2.5, 2.0, 1.8, 1.8, 2.0, 2.0]
)

heading(doc, "질문 타입별 Judge Score (QASPER)", level=2)
add_table(doc,
    ["Baseline", "local_factual", "global_synthesis", "terminology_sensitive"],
    [
        ["Direct QA",           "0.567 ★", "0.400", "0.367"],
        ["Summary-Mediated QA", "0.467",   "0.400", "0.333"],
        ["Standard RAG",        "0.433",   "0.333", "0.200 ▼"],
    ],
    col_widths=[4.0, 3.5, 3.8, 4.2]
)

body(doc, "▶ 핵심 패턴: Human-annotated 질문에서 전반적으로 점수 하락. "
         "Direct QA가 QASPER에서 1위 (Judge=0.444) — 문서 전체를 보는 방식이 "
         "복잡한 NLP 논문 질문에 유리.", indent=True)

doc.add_paragraph()
heading(doc, "3-3. Token F1 vs. Judge Score 비교 (핵심 발견)", level=2)
add_table(doc,
    ["데이터셋", "Baseline", "Token F1", "Judge Score", "차이 (Judge − F1)"],
    [
        ["PubMedQA", "Direct QA",           "0.497", "0.844", "+0.347"],
        ["PubMedQA", "Summary-Mediated QA", "0.526", "0.878", "+0.352"],
        ["PubMedQA", "Standard RAG",        "0.364", "0.622", "+0.258"],
        ["QASPER",   "Direct QA",           "0.240", "0.444", "+0.204"],
        ["QASPER",   "Summary-Mediated QA", "0.222", "0.400", "+0.178"],
    ],
    col_widths=[2.5, 4.0, 2.5, 2.8, 3.5]
)
body(doc, "▶ Token F1은 실제 성능을 평균 +0.27p 과소평가. "
         "모델이 의미적으로 맞는 답을 했음에도 표현 차이로 낮은 점수 → "
         "단일 지표 의존 시 시스템 성능 오판 가능.", indent=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 4. Limitation 및 개선 가능성
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "4. Limitation 및 개선 가능성")

add_table(doc,
    ["#", "Limitation", "심각도", "개선 가능?", "방법"],
    [
        ["L1", "데이터 규모 부족 (45개)",              "높음", "✅ 즉시 가능",  "QASPER 파라미터 변경으로 150개 확장"],
        ["L2", "PubMedQA 질문이 GPT 생성",             "높음", "✅ 이미 해결",  "QASPER(human)으로 대체 완료"],
        ["L3", "Token F1 과소평가",                    "높음", "✅ 즉시 가능",  "BERTScore 추가, LLM Judge 이미 구현"],
        ["L4", "LightRAG를 단일 문서에 사용",          "중간", "🟡 부분 가능",  "멀티 문서 실험 설계 변경 또는 조건 명시"],
        ["L5", "LightRAG QASPER 미실행",               "중간", "✅ 즉시 가능",  "명령어 한 줄로 실행 가능"],
        ["L6", "Chunk 파라미터 고정 (미검증)",         "중간", "✅ 가능",       "Ablation 스크립트로 자동화"],
        ["L7", "Judge도 GPT (순환 편향)",              "낮음", "🟡 부분 가능",  "다른 LLM Judge 또는 Human eval 추가"],
        ["L8", "Hit Rate가 Binary (세밀도 부족)",      "낮음", "✅ 즉시 가능",  "Recall@k / MRR로 교체"],
        ["L9", "영어 데이터만 사용",                   "낮음", "❌ 구조적 한계","별도 프로젝트 필요"],
    ],
    col_widths=[0.7, 4.0, 1.8, 2.3, 5.0],
    header_color="365F91"
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 5. 주요 발견 및 결론
# ══════════════════════════════════════════════════════════════
heading(doc, "5. 주요 발견 및 결론")

heading(doc, "5-1. 어떤 방식이 가장 좋은가?", level=2)
bullet(doc, "전반적 최고 성능: Summary-Mediated QA (PubMedQA Judge=0.878, incorrect 0건)", bold_prefix="✅")
bullet(doc, "단순성 대비 최고 성능: Direct QA (QASPER에서 3방식 중 1위)", bold_prefix="✅")
bullet(doc, "LightRAG는 단일 문서 조건에서 부적합 — 멀티 문서 환경에서 재평가 필요", bold_prefix="⚠️")

heading(doc, "5-2. 평가 지표에 대한 발견", level=2)
bullet(doc, "Token F1만으로는 RAG 시스템을 공정하게 평가할 수 없음 (평균 +0.27p 과소평가)", bold_prefix="🔑")
bullet(doc, "LLM-as-judge와 Token F1을 함께 사용하는 이중 평가 체계가 필요", bold_prefix="🔑")
bullet(doc, "Human-annotated 질문(QASPER)이 GPT 생성 질문보다 현실적 난이도 반영", bold_prefix="🔑")

heading(doc, "5-3. Standard RAG의 취약점", level=2)
bullet(doc, "Retrieval Hit=1.00임에도 terminology_sensitive Judge=0.200 — "
           "청크를 찾았지만 전문 용어 표현 차이로 답변 실패", bold_prefix="📌")
bullet(doc, "검색은 성공해도 답변이 틀리는 'retrieval ≠ correctness' 현상 확인", bold_prefix="📌")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 6. 향후 계획
# ══════════════════════════════════════════════════════════════
heading(doc, "6. 향후 계획 (우선순위 순)")
add_table(doc,
    ["우선순위", "작업", "예상 소요", "기대 효과"],
    [
        ["1순위", "LightRAG → QASPER 실행",               "~2시간",   "4-way 공정 비교 완성"],
        ["2순위", "데이터 규모 확장 (45 → 150개)",         "~반나절",  "통계적 신뢰성 확보"],
        ["3순위", "BERTScore 추가",                        "~2시간",   "의미 기반 평가 다각화"],
        ["4순위", "Chunk size / top-k ablation",           "~반나절",  "Standard RAG 최적 파라미터 도출"],
        ["5순위", "보고서/논문 초안 작성",                  "~1일",     "실험 결과 공식화"],
    ],
    col_widths=[2.0, 5.0, 2.5, 5.0]
)

doc.add_paragraph()
note_p = doc.add_paragraph()
r = note_p.add_run("※ 생성된 모든 시각화 파일 위치: ")
set_font(r, size=9.5, italic=True, color=(89,89,89))
r2 = note_p.add_run("results/analysis/ 폴더")
set_font(r2, size=9.5, bold=True, italic=True, color=(31,73,125))
r3 = note_p.add_run("  |  실험 코드: ")
set_font(r3, size=9.5, italic=True, color=(89,89,89))
r4 = note_p.add_run("/Users/seungwookim/Desktop/rag_experiment/")
set_font(r4, size=9.5, bold=True, italic=True, color=(31,73,125))

# ── 저장 ────────────────────────────────────────────────────
out = "/Users/seungwookim/Desktop/rag_experiment/RAG_실험_보고서.docx"
doc.save(out)
print(f"저장 완료: {out}")
